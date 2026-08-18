#!/usr/bin/env python3
"""Generate Dispatcher Staffing Assessment — HTML + DOCX from analysis JSON."""
from __future__ import annotations

import json
import os
from datetime import date

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
JSON_PATH = os.path.join(ROOT, "reports", "dispatcher_fleet_analysis.json")
OUT_HTML = os.path.join(ROOT, "reports", "Dispatcher_Staffing_Assessment_2026.html")
OUT_DOCX = os.path.join(ROOT, "reports", "Dispatcher_Staffing_Assessment_2026.docx")

# Lean / strong-FMS benchmark (trucks supervised per dispatcher per shift)
LEAN_RATIO_LO = 60
LEAN_RATIO_HI = 80
TYPICAL_RATIO = 50
HEAVY_RATIO = 40

MAIN_CONTRACTORS = ("RIM", "SSS", "PPP", "SMA", "STM", "CKB", "GMG", "HJS")


def load_data():
    with open(JSON_PATH, encoding="utf-8") as f:
        return json.load(f)


def pick_contractors(data):
    rows = []
    for r in data["contractor_summary"]:
        co = r["contractor"]
        if co in MAIN_CONTRACTORS or (r.get("jan_apr_avg_running_trucks") or 0) >= 30:
            if "?" in co:
                continue
            rows.append(r)
    return rows


def dispatch_range(trucks, lo=LEAN_RATIO_LO, hi=LEAN_RATIO_HI):
    if not trucks:
        return "—", "—"
    import math
    return str(math.ceil(trucks / hi)), str(math.ceil(trucks / lo))


def site_totals(rows):
    avg = sum(r.get("jan_apr_avg_running_trucks") or 0 for r in rows if r.get("jan_apr_avg_running_trucks"))
    # add minor contractors from json not in main list
    with open(JSON_PATH, encoding="utf-8") as f:
        all_rows = json.load(f)["contractor_summary"]
    site_avg = sum(r.get("jan_apr_avg_running_trucks") or 0 for r in all_rows if r.get("jan_apr_avg_running_trucks") and "?" not in r["contractor"])
    p90_vals = [r.get("jan_apr_p90_running_trucks") for r in all_rows if r.get("jan_apr_p90_running_trucks") and "?" not in r["contractor"]]
    max_vals = [r.get("jan_apr_max_running_trucks") for r in all_rows if r.get("jan_apr_max_running_trucks") and "?" not in r["contractor"]]
    return {
        "site_avg": round(site_avg, 0),
        "site_p90": round(max(p90_vals) if p90_vals else 0, 0),  # approximate from components
        "site_max_day": 1685,
        "measured_avg": 1153,
        "measured_median": 1140,
        "measured_p90": 1443,
    }


def build_html(data, contractors, site):
    today = date.today().strftime("%d %B %Y")
    lean_lo, lean_hi = dispatch_range(site["measured_avg"])
    lean_p90_lo, lean_p90_hi = dispatch_range(site["measured_p90"])
    lean_max_lo, lean_max_hi = dispatch_range(site["site_max_day"])

    rows_html = ""
    for r in contractors:
        co = r["contractor"]
        run = r.get("jan_apr_avg_running_trucks")
        p90 = r.get("jan_apr_p90_running_trucks")
        plan = r.get("jan_apr_avg_planned_dt")
        reg_se = r.get("safety_enviro_dump_trucks") or 0
        reg_wbn = r.get("wbn_register_dt") or 0
        d_lo, d_hi = dispatch_range(run)
        dp90_lo, dp90_hi = dispatch_range(p90)
        rows_html += f"""<tr>
          <td><strong>{co}</strong></td>
          <td class="num">{reg_se:,}</td>
          <td class="num">{reg_wbn:,}</td>
          <td class="num">{plan if plan else '—'}</td>
          <td class="num">{run if run else '—'}</td>
          <td class="num">{p90 if p90 else '—'}</td>
          <td class="num">{d_lo}–{d_hi}</td>
          <td class="num">{dp90_lo}–{dp90_hi}</td>
        </tr>"""

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1"/>
<title>Dispatcher Staffing Assessment — IWIP Haul Fleet</title>
<style>
  :root {{
    --ink: #1a2332;
    --muted: #5c6b7a;
    --accent: #0d4f8b;
    --accent-light: #e8f2fb;
    --border: #d0dae4;
    --bg: #f7f9fc;
  }}
  * {{ box-sizing: border-box; }}
  body {{
    font-family: "Segoe UI", Calibri, Arial, sans-serif;
    color: var(--ink);
    line-height: 1.55;
    max-width: 920px;
    margin: 0 auto;
    padding: 48px 32px 64px;
    background: #fff;
  }}
  .cover {{
    border-bottom: 4px solid var(--accent);
    padding-bottom: 28px;
    margin-bottom: 36px;
  }}
  .cover h1 {{
    font-size: 1.75rem;
    font-weight: 700;
    color: var(--accent);
    margin: 0 0 8px;
    letter-spacing: -0.02em;
  }}
  .cover .subtitle {{ font-size: 1.1rem; color: var(--muted); margin: 0 0 20px; }}
  .meta {{ font-size: 0.9rem; color: var(--muted); }}
  .meta span {{ display: inline-block; margin-right: 24px; }}
  h2 {{
    font-size: 1.15rem;
    color: var(--accent);
    border-bottom: 1px solid var(--border);
    padding-bottom: 6px;
    margin-top: 32px;
  }}
  h3 {{ font-size: 1rem; margin-top: 20px; color: var(--ink); }}
  p {{ margin: 0.75em 0; }}
  .callout {{
    background: var(--accent-light);
    border-left: 4px solid var(--accent);
    padding: 16px 20px;
    margin: 20px 0;
    font-size: 0.95rem;
  }}
  .callout strong {{ color: var(--accent); }}
  table {{
    width: 100%;
    border-collapse: collapse;
    font-size: 0.88rem;
    margin: 16px 0;
  }}
  th, td {{
    border: 1px solid var(--border);
    padding: 8px 10px;
    text-align: left;
  }}
  th {{
    background: var(--accent);
    color: #fff;
    font-weight: 600;
  }}
  tr:nth-child(even) {{ background: var(--bg); }}
  td.num {{ text-align: right; font-variant-numeric: tabular-nums; }}
  ul {{ margin: 0.5em 0; padding-left: 1.4em; }}
  li {{ margin: 0.35em 0; }}
  .footnote {{
    font-size: 0.82rem;
    color: var(--muted);
    margin-top: 40px;
    border-top: 1px solid var(--border);
    padding-top: 16px;
  }}
  @media print {{
    body {{ padding: 24px; }}
    h2 {{ page-break-after: avoid; }}
    table {{ page-break-inside: avoid; }}
  }}
</style>
</head>
<body>

<div class="cover">
  <h1>Dispatcher Staffing Assessment</h1>
  <p class="subtitle">IWIP Haul Fleet — Control Room Manning vs Registered Fleet &amp; Actual Operations</p>
  <div class="meta">
    <span><strong>Period analysed:</strong> 1 January – 30 April 2026</span>
    <span><strong>Prepared:</strong> {today}</span>
    <span><strong>Classification:</strong> Internal — Operations Planning</span>
  </div>
</div>

<h2>1. Executive Summary</h2>
<p>This assessment compares <strong>registered fleet</strong> (Safety &amp; Environment equipment register and WBN operations register) against <strong>actual haul activity</strong> (weighbridge tickets) and <strong>daily mining plans</strong> for January–April 2026. The objective is to support control-room dispatcher staffing by contractor and site-wide.</p>

<div class="callout">
  <strong>Lean / strong FMS benchmark:</strong> International practice for operations with a mature Fleet Management System (automated assignment, live GPS, cycle-time dashboards) supports approximately <strong>{LEAN_RATIO_LO}–{LEAN_RATIO_HI} active haul trucks per dispatcher per shift</strong>. At the measured site average of <strong>{site["measured_avg"]:,.0f} trucks/day</strong>, this implies <strong>{lean_lo}–{lean_hi} dispatchers per shift</strong> for integrated control-room coverage. On busy days (P90 ≈ {site["measured_p90"]:,.0f} trucks), plan for <strong>{lean_p90_lo}–{lean_p90_hi} dispatchers per shift</strong>; peak observed day ({site["site_max_day"]:,.0f} trucks) requires <strong>{lean_max_lo}–{lean_max_hi} dispatchers per shift</strong>.
</div>

<p><strong>Key findings:</strong></p>
<ul>
  <li>Registered dump trucks across main contractors total <strong>~2,500+ units</strong> on books; only <strong>~1,150 trucks/day</strong> actually haul on IWIP roads (average Jan–Apr).</li>
  <li><strong>RIM</strong> dominates activity (~318 trucks/day average, up to 472 on peak days) — largest dispatcher requirement.</li>
  <li>Daily plans (<code>EQUIPMENTS_PLAN</code>) align reasonably with actual running fleet for major contractors; use <em>running trucks</em>, not register totals, for manning.</li>
  <li>No fixed ISO or global standard exists for “dispatchers per truck”; lean FMS sites benchmark by <strong>supervised truck count</strong>, pit complexity, and shift structure (typically two 12-hour shifts).</li>
</ul>

<h2>2. Scope &amp; Data Sources</h2>
<table>
  <tr><th>Source</th><th>Database / Table</th><th>Purpose</th></tr>
  <tr><td>Safety equipment register</td><td><code>SAFETY_ENVIRO.dbo.EQUIPMENTS</code></td><td>Total fleet by Company and Equipment_Type (5,234 records)</td></tr>
  <tr><td>Operations register</td><td><code>WBN_DATABASE.dbo.EQUIPMENTS</code></td><td>Contractor fleet by TYPE / ID_EQ (7,221 records)</td></tr>
  <tr><td>Daily mining plan</td><td><code>WBN_DATABASE.dbo.EQUIPMENTS_PLAN</code></td><td>Planned DT/ADT/EXCA units per contractor per day</td></tr>
  <tr><td>Actual haul proof</td><td><code>WBN_DATABASE.dbo.HAULAGE_IWIP_CLEAN</code></td><td>Unique TRUCK_ID per contractor per day (Jan–Apr 2026)</td></tr>
</table>

<h2>3. Registered Fleet — Main Haul Contractors</h2>
<p>Counts from <strong>SAFETY_ENVIRO</strong> (Company / Equipment_Type). “Dump truck” includes rigid dump truck, ADT, and EV variants.</p>
<table>
  <tr><th>Contractor</th><th>Total units</th><th>Dump trucks</th><th>Excavators</th><th>Notes</th></tr>
  <tr><td>RIM</td><td class="num">2,683</td><td class="num">1,407</td><td class="num">331</td><td>828 DT + 435 EV + 144 ADT</td></tr>
  <tr><td>SMA</td><td class="num">807</td><td class="num">375</td><td class="num">108</td><td>297 DT + 78 ADT</td></tr>
  <tr><td>PPP</td><td class="num">475</td><td class="num">265</td><td class="num">72</td><td>229 DT + 36 ADT</td></tr>
  <tr><td>SSS</td><td class="num">253</td><td class="num">178</td><td class="num">15</td><td>—</td></tr>
  <tr><td>STM</td><td class="num">248</td><td class="num">98</td><td class="num">53</td><td>55 DT + 43 ADT</td></tr>
  <tr><td>CKB</td><td class="num">127</td><td class="num">105</td><td class="num">5</td><td>—</td></tr>
  <tr><td>GMG</td><td class="num">117</td><td class="num">78</td><td class="num">15</td><td>—</td></tr>
</table>

<h2>4. Jan–Apr 2026 — Planned vs Actual Running Trucks</h2>
<p>“Running” = distinct truck IDs with ≥1 weighbridge ticket that day. This is the correct basis for dispatcher load, not register totals.</p>
<table>
  <tr>
    <th>Contractor</th>
    <th>Registered DT (Safety)</th>
    <th>Registered DT (WBN)</th>
    <th>Avg planned DT/day</th>
    <th>Avg running/day</th>
    <th>P90 running/day</th>
    <th>Dispatchers/shift<br/><small>lean {LEAN_RATIO_LO}–{LEAN_RATIO_HI}:1</small></th>
    <th>Dispatchers/shift<br/><small>peak P90</small></th>
  </tr>
  {rows_html}
</table>

<p><strong>Site total (all contractors, Jan–Apr):</strong> average <strong>{site["measured_avg"]:,.0f}</strong> trucks/day · median {site["measured_median"]:,.0f} · P90 {site["measured_p90"]:,.0f} · maximum observed {site["site_max_day"]:,.0f}.</p>

<h2>5. International Benchmarks — Lean / Strong FMS Operations</h2>
<p>There is <strong>no universal ISO ratio</strong> for dispatchers per truck. Industry guidance (AusIMM Mine Monitoring &amp; Control, major FMS vendors, open-pit dispatch literature) converges on the following:</p>

<h3>5.1 Fleet Management System maturity</h3>
<ul>
  <li><strong>Strong FMS:</strong> Real-time GPS, automated truck assignment, cycle/queue dashboards, shift handover tools — reduces manual reassignment and allows higher supervised fleet per dispatcher.</li>
  <li><strong>Lean control room:</strong> Primary dispatcher focuses on exceptions (breakdowns, road closures, shovel moves); system handles routine assignment.</li>
  <li><strong>Match factor</strong> (truck arrival rate vs shovel service rate) drives productivity; staffing follows <em>supervised active fleet</em> and pit complexity, not register size.</li>
</ul>

<h3>5.2 Benchmark ratios (haul trucks per dispatcher per shift)</h3>
<table>
  <tr><th>Operating model</th><th>Trucks per dispatcher</th><th>Typical use</th></tr>
  <tr><td><strong>Lean / strong FMS</strong></td><td class="num"><strong>{LEAN_RATIO_LO}–{LEAN_RATIO_HI}</strong></td><td>Mature FMS, stable routes, locked shovel assignments, automated logging</td></tr>
  <tr><td>Typical multi-contractor pit</td><td class="num">{TYPICAL_RATIO}–{LEAN_RATIO_LO}</td><td>Multiple load areas (TF, BLB, KR), mixed contractors — <em>current IWIP profile</em></td></tr>
  <tr><td>Heavy manual coordination</td><td class="num">{HEAVY_RATIO}–{TYPICAL_RATIO}</td><td>Frequent roadworks, weak FMS, high breakdown rate</td></tr>
</table>

<h3>5.3 Control room structure (international practice)</h3>
<ul>
  <li><strong>Two shifts per day</strong> (12 h) for 24-hour coverage; separate day and night rosters.</li>
  <li><strong>One senior controller / shift supervisor</strong> per shift with decision authority (AusIMM MMC Part 3).</li>
  <li><strong>Relief depth:</strong> ~1 trainee or relief dispatcher per 4–5 primary dispatchers.</li>
  <li>Contractors may maintain <strong>dedicated dispatch desks</strong> within a shared IWIP control room, or a central pool — organisational choice; truck counts still apply per supervised fleet.</li>
</ul>

<h2>6. Recommended Dispatcher Manning</h2>

<h3>6.1 Site-wide (integrated control room) — lean FMS target ({LEAN_RATIO_LO}–{LEAN_RATIO_HI}:1)</h3>
<table>
  <tr><th>Scenario</th><th>Active haul trucks</th><th>Dispatchers per shift</th></tr>
  <tr><td>Average day (Jan–Apr measured)</td><td class="num">{site["measured_avg"]:,.0f}</td><td class="num"><strong>{lean_lo}–{lean_hi}</strong></td></tr>
  <tr><td>Busy day (P90)</td><td class="num">{site["measured_p90"]:,.0f}</td><td class="num"><strong>{lean_p90_lo}–{lean_p90_hi}</strong></td></tr>
  <tr><td>Peak observed day</td><td class="num">{site["site_max_day"]:,.0f}</td><td class="num"><strong>{lean_max_lo}–{lean_max_hi}</strong></td></tr>
</table>
<p>Add <strong>1 shift supervisor</strong> per shift. For 24 h operations: multiply dispatcher headcount by <strong>2 shifts</strong> (e.g. average-day roster ≈ {(int(lean_lo)+int(lean_hi))}–{(int(lean_lo)+int(lean_hi))*2} dispatcher slots across day + night, plus 2 supervisors).</p>

<h3>6.2 By contractor (dedicated desks) — lean FMS {LEAN_RATIO_LO}–{LEAN_RATIO_HI}:1</h3>
<p>If each contractor staffs its own dispatch position(s), use average running trucks and peak (P90) from Section 4. Largest requirements: <strong>RIM 5–8</strong>, <strong>SSS 2–3</strong>, <strong>PPP 1–2</strong>, <strong>SMA 1–2</strong>, <strong>STM 1–2</strong>, <strong>CKB 1</strong>, <strong>GMG 1</strong> dispatchers per shift at the lean benchmark.</p>

<h3>6.3 Planning rule (recommended)</h3>
<div class="callout">
  Staff dispatch using: <strong>Dispatchers per shift = ⌈Active haul trucks ÷ 60⌉ to ⌈Active haul trucks ÷ 80⌉</strong> under lean FMS, plus one supervisor. Use <strong>P90 daily truck count</strong> for roster sizing, not register totals or average-only (under-staffs peak season).
</div>

<h2>7. Limitations &amp; Caveats</h2>
<ul>
  <li>Register counts include maintenance, standby, and off-site units — not dispatch workload.</li>
  <li>Weighbridge counts only material hauls; road-only / HRM trucks may not appear.</li>
  <li>Some contractor codes in haulage data are encoding-corrupted; ~50 trucks/day may sit under unresolved labels.</li>
  <li>HJS has registered fleet but negligible Jan–Apr IWIP haul activity.</li>
  <li>Dispatcher need also depends on excavator count, road segments, and FMS maturity — ratios should be validated after 90 days of control-room KPI tracking.</li>
</ul>

<h2>8. References</h2>
<ul>
  <li>AusIMM Bulletin — <em>Mine monitoring or control (Part 3): Staffing the control room</em></li>
  <li>GroundHog / industry FMS — Dispatch Operator Handbook (fleet visibility, automated cycle logging)</li>
  <li>Open-pit dispatch literature — match factor &amp; shovel–truck balance ( Douglas, Burt &amp; Caccetta; E3S Conf. functional quality criterion)</li>
  <li>Internal data — WBN FMS Simulator analysis, August 2026</li>
</ul>

<p class="footnote">
  Generated from live database queries against SAFETY_ENVIRO and WBN_DATABASE. 
  Analysis script: <code>scripts/generate_dispatcher_report.py</code> · 
  Data cache: <code>reports/dispatcher_fleet_analysis.json</code>
</p>

</body>
</html>"""


def build_docx(data, contractors, site):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import math

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    t = doc.add_heading("Dispatcher Staffing Assessment", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("IWIP Haul Fleet — Control Room Manning vs Registered Fleet & Actual Operations")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Period: 1 January – 30 April 2026  |  Prepared: {date.today():%d %B %Y}  |  Internal — Operations Planning")
    doc.add_paragraph()

    lean_lo = math.ceil(site["measured_avg"] / LEAN_RATIO_HI)
    lean_hi = math.ceil(site["measured_avg"] / LEAN_RATIO_LO)
    lean_p90_lo = math.ceil(site["measured_p90"] / LEAN_RATIO_HI)
    lean_p90_hi = math.ceil(site["measured_p90"] / LEAN_RATIO_LO)

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This assessment compares registered fleet against actual haul activity and daily mining plans "
        "for January–April 2026 to support control-room dispatcher staffing by contractor and site-wide."
    )
    p = doc.add_paragraph()
    p.add_run("Lean / strong FMS benchmark: ").bold = True
    p.add_run(
        f"Operations with a mature Fleet Management System support approximately "
        f"{LEAN_RATIO_LO}–{LEAN_RATIO_HI} active haul trucks per dispatcher per shift. "
        f"At the measured site average of {site['measured_avg']:,.0f} trucks/day, this implies "
        f"{lean_lo}–{lean_hi} dispatchers per shift. Busy days (P90 ≈ {site['measured_p90']:,.0f} trucks) "
        f"require {lean_p90_lo}–{lean_p90_hi} dispatchers per shift."
    )

    doc.add_heading("2. Key Findings", level=1)
    for item in [
        "Registered dump trucks (~2,500+ main contractors) far exceed daily running fleet (~1,150 trucks/day average).",
        "RIM dominates (~318 trucks/day average, peak 472) — largest dispatcher requirement.",
        "Use running trucks from weighbridge, not register totals, for manning decisions.",
        "No fixed ISO ratio exists; lean FMS sites use 60–80 trucks per dispatcher per shift.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Registered Fleet (SAFETY_ENVIRO)", level=1)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Contractor", "Total units", "Dump trucks", "Excavators", "Notes"]):
        hdr[i].text = h
    reg_rows = [
        ("RIM", "2,683", "1,407", "331", "828 DT + 435 EV + 144 ADT"),
        ("SMA", "807", "375", "108", "297 DT + 78 ADT"),
        ("PPP", "475", "265", "72", "229 DT + 36 ADT"),
        ("SSS", "253", "178", "15", "—"),
        ("STM", "248", "98", "53", "55 DT + 43 ADT"),
        ("CKB", "127", "105", "5", "—"),
        ("GMG", "117", "78", "15", "—"),
    ]
    for row in reg_rows:
        cells = tbl.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    doc.add_heading("4. Jan–Apr 2026 — Planned vs Actual Running", level=1)
    tbl2 = doc.add_table(rows=1, cols=8)
    tbl2.style = "Table Grid"
    h2 = ["Contractor", "Reg DT (Safety)", "Reg DT (WBN)", "Avg plan/day", "Avg run/day", "P90 run/day",
          f"Disp/shift ({LEAN_RATIO_LO}–{LEAN_RATIO_HI}:1)", "Disp/shift P90"]
    for i, h in enumerate(h2):
        tbl2.rows[0].cells[i].text = h
    for r in contractors:
        if "?" in r["contractor"]:
            continue
        run = r.get("jan_apr_avg_running_trucks")
        p90 = r.get("jan_apr_p90_running_trucks")
        d_lo, d_hi = dispatch_range(run)
        dp_lo, dp_hi = dispatch_range(p90)
        cells = tbl2.add_row().cells
        cells[0].text = r["contractor"]
        cells[1].text = f"{r.get('safety_enviro_dump_trucks') or 0:,}"
        cells[2].text = f"{r.get('wbn_register_dt') or 0:,}"
        cells[3].text = str(r.get("jan_apr_avg_planned_dt") or "—")
        cells[4].text = str(run or "—")
        cells[5].text = str(p90 or "—")
        cells[6].text = f"{d_lo}–{d_hi}"
        cells[7].text = f"{dp_lo}–{dp_hi}"

    doc.add_paragraph(
        f"Site total: average {site['measured_avg']:,.0f} trucks/day · P90 {site['measured_p90']:,.0f} · "
        f"peak {site['site_max_day']:,.0f}."
    )

    doc.add_heading("5. International Benchmarks — Lean / Strong FMS", level=1)
    doc.add_paragraph(
        "No universal ISO standard exists for dispatchers per truck. Industry practice for mature FMS operations:"
    )
    tbl3 = doc.add_table(rows=1, cols=3)
    tbl3.style = "Table Grid"
    for i, h in enumerate(["Model", "Trucks per dispatcher", "Application"]):
        tbl3.rows[0].cells[i].text = h
    bench = [
        ("Lean / strong FMS", f"{LEAN_RATIO_LO}–{LEAN_RATIO_HI}", "Mature FMS, automated assignment, stable routes"),
        ("Typical multi-contractor", f"{TYPICAL_RATIO}–{LEAN_RATIO_LO}", "TF/BLB/KR mixed pits — IWIP profile"),
        ("Heavy manual", f"{HEAVY_RATIO}–{TYPICAL_RATIO}", "Roadworks, weak FMS, high breakdowns"),
    ]
    for row in bench:
        cells = tbl3.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    doc.add_heading("6. Recommended Manning (Lean FMS 60–80:1)", level=1)
    tbl4 = doc.add_table(rows=1, cols=3)
    tbl4.style = "Table Grid"
    for i, h in enumerate(["Scenario", "Active trucks", "Dispatchers/shift"]):
        tbl4.rows[0].cells[i].text = h
    for scen, trucks in [
        ("Average day", site["measured_avg"]),
        ("Busy day (P90)", site["measured_p90"]),
        ("Peak observed", site["site_max_day"]),
    ]:
        lo, hi = dispatch_range(trucks)
        cells = tbl4.add_row().cells
        cells[0].text = scen
        cells[1].text = f"{trucks:,.0f}"
        cells[2].text = f"{lo}–{hi}"

    doc.add_paragraph(
        "Add 1 shift supervisor per shift. Two 12-hour shifts for 24 h coverage. "
        "Staff using P90 truck counts for roster sizing, not register totals."
    )

    doc.add_heading("7. Limitations", level=1)
    for item in [
        "Register includes off-site and non-haul units.",
        "Weighbridge = material haul only.",
        "HJS registered but inactive Jan–Apr on IWIP haul.",
        "Validate ratios with 90 days of control-room KPIs.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("8. References", level=1)
    refs = [
        "AusIMM — Mine monitoring or control (Part 3): Staffing the control room",
        "GroundHog FMS — Dispatch Operator Handbook",
        "Open-pit dispatch / match factor literature",
        "WBN FMS Simulator internal analysis, August 2026",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Bullet")

    doc.save(OUT_DOCX)


def main():
    data = load_data()
    contractors = pick_contractors(data)
    site = site_totals(contractors)

    html = build_html(data, contractors, site)
    os.makedirs(os.path.dirname(OUT_HTML), exist_ok=True)
    with open(OUT_HTML, "w", encoding="utf-8") as f:
        f.write(html)
    print("Wrote", OUT_HTML)

    build_docx(data, contractors, site)
    print("Wrote", OUT_DOCX)


if __name__ == "__main__":
    main()
